#!/usr/bin/env python3
"""Build manuscript-styled Word appendices and a tracked-correction copy."""

from __future__ import annotations

import argparse
import copy
import datetime as dt
import hashlib
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
XML_NS = "http://www.w3.org/XML/1998/namespace"
NS = {"w": W_NS, "m": M_NS}


SOURCE_PARAGRAPHS = {
    "model_intro": (
        "To analyze excess mortality in each stratified population across different "
        "regions, we adopt the following Bayesian hierarchical model:"
    ),
    "model_explanation": (
        "Here y(tᵢ) denotes the all-cause mortality count for the given "
        "stratification group in the specified region at time tᵢ (measured in "
        "years), and n denotes the number of observations taken prior to January "
        "1, 2020. The model includes three main components: a long-term trend "
        "fₜᵣ, which follows a second-order integrated Wiener process (Zhang et "
        "al., 2024); a seasonal component fₛ, which follows a yearly seasonal "
        "Gaussian process with four harmonics (Zhang et al., 2023); and an "
        "observational-level random error term εᵢ to account for "
        "overdispersion. Independent Exponential priors, motivated by Simpson et "
        "al. (2017), are used for the variance parameters σₜᵣ, σₛ, and "
        "σₑ. Additional details on prior specification and Bayesian "
        "computation are provided in the supplementary material."
    ),
    "us_offset": (
        "For the United States, mortality data are available at the monthly "
        "level, so y(tᵢ) represents monthly all-cause mortality counts. To "
        "adjust for differences in the number of days per month, an offset equal "
        "to the month length was included in the model."
    ),
    "prediction": (
        "The model in Eq. (1) was fit using pre-pandemic mortality data (before "
        "January 1, 2020). Forecasts of expected mortality were obtained as "
        "posterior predictive samples. Excess mortality at time t was defined as "
        "the difference between the observed mortality y(t) and the forecasted "
        "mortality ŷ(t). To account for population structure, we adopt the "
        "P-score metric from Knutson et al. (2023), which normalizes excess deaths "
        "by forecasted mortality:"
    ),
    "combined_intro": (
        "In a given region, we also define the all-age P-score pᵃˡˡ(t) as the "
        "P-score calculated from the aggregated mortality counts of two (or "
        "three) age groups of interest (0-44, 45-64, and 65-84 for the US "
        "age-sex-stratified analysis; 40-59 and 60-79 for other analyses):"
    ),
    "combined_where": (
        "where yʲ(t) and ŷʲ(t) denote the observed and predicted death "
        "counts for age group j."
    ),
    "pooled_intro": (
        "When comparing P-score series across groups of regions, pooled estimates "
        "were obtained through inverse-variance weighting:"
    ),
    "pooled_where": (
        "where πₖ denotes the normalized inverse variance of the P-score in "
        "region k."
    ),
    "b_definition": (
        "The P-score at time t is defined as in Eq. (2), with observed mortality "
        "y(t) and predicted baseline mortality ŷ(t)."
    ),
    "b_combined": (
        "For age-aggregated analyses, we calculated P-scores for all ages combined "
        "by adding up deaths across age groups before computing the percentage. "
        "This means age groups with more deaths have more influence on the "
        "combined result. For selected analyses, all-age P-scores aggregated "
        "across the 40-59 and 60-79 year age groups were computed."
    ),
    "b_pooled": (
        "When combining P-scores from multiple regions (for example, comparing "
        "high- versus low-vaccination countries), pooled P-scores were calculated "
        "via inverse-variance weighting, giving more weight to regions with more "
        "reliable estimates. This ensures that areas with better data quality "
        "have greater influence on the overall comparison."
    ),
}


CORRECTED_PARAGRAPHS = {
    "model_intro": (
        "To analyze excess mortality in each stratified population across different "
        "regions, we adopt the following Bayesian log-Poisson time-series model, "
        "fitted independently to each region-by-stratum series:"
    ),
    "model_explanation": (
        "Here y(tᵢ) denotes the all-cause mortality count for the given "
        "stratification group in the specified region at time tᵢ (measured in "
        "years), n denotes the number of observations taken prior to January 1, "
        "2020, β₀ is an intercept, and oᵢ is a known offset. The model "
        "includes three main components: a long-term trend fₜᵣ, which follows a "
        "second-order integrated Wiener process (Zhang et al., 2024); a seasonal "
        "component fₛ, which follows a yearly seasonal Gaussian process with four "
        "harmonics (Zhang et al., 2023); and an observation-level Gaussian random "
        "effect εᵢ to account for overdispersion. Independent "
        "penalized-complexity priors, motivated by Simpson et al. (2017), are "
        "placed on the standard-deviation scales. The trend prior is calibrated "
        "by P(five-year predictive SD > 0.1) = 0.01, the seasonal prior by "
        "P(one-year predictive SD > 0.1) = 0.01, and the overdispersion prior "
        "by P(σₑ > 0.1) = "
        "0.01."
    ),
    "us_offset": (
        "For the United States, mortality data are available at the monthly "
        "level, so y(tᵢ) represents monthly all-cause mortality counts. To "
        "adjust for differences in the number of days per month, oᵢ is the log "
        "of the number of calendar days in that month; oᵢ is zero for the "
        "weekly analyses."
    ),
    "combined_intro": (
        "In a given region, we also define the combined-age P-score as the P-score "
        "calculated after aggregating the mortality counts across "
        "the age groups of interest (0-44, 45-64, and 65-84 for the US "
        "age-sex-stratified analysis, giving ages 0-84; 40-59 and 60-79 for the "
        "other combined-age analyses, giving ages 40-79):"
    ),
    "b_combined": (
        "For age-aggregated analyses, we calculated P-scores for the selected "
        "combined-age range by adding deaths across age groups before computing "
        "the percentage. This means age groups with more deaths have more "
        "influence on the combined result. For selected analyses, combined-age "
        "P-scores across the 40-59 and 60-79 year age groups were computed, "
        "corresponding to ages 40-79 rather than all ages."
    ),
    "combined_where": (
        "where yʲ(t) and ŷʲ(t) denote the observed and predicted death counts "
        "for age group j, and J is the number of included age groups (J = 3 for "
        "US ages 0-84 and J = 2 for ages 40-79)."
    ),
    "b_pooled": (
        "When combining P-scores from multiple regions (for example, comparing "
        "high- versus low-vaccination countries), pooled P-scores were calculated "
        "via inverse-variance weighting, giving more weight to regions with "
        "smaller estimated posterior variances."
    ),
}


BASELINE_FORMULAS = {
    "eq_model": r"\begin{aligned} y(t_i) &\overset{\mathrm{ind}}{\sim} "
    r"\operatorname{Poisson}(\lambda(t_i)),\quad \forall i \in [n] \\ "
    r"\lambda(t_i) &= f_{\mathrm{tr}}(t_i)+f_s(t_i)+\epsilon_i, \\ "
    r"f_{\mathrm{tr}} &\sim \operatorname{IWP}_2(\sigma_{\mathrm{tr}}),\quad "
    r"f_s \sim \operatorname{sGP}_1(\sigma_s),\quad "
    r"\epsilon_i \overset{\mathrm{iid}}{\sim} N(0,\sigma_\epsilon). "
    r"\end{aligned}",
    "eq_pscore": r"p(t)=\frac{y(t)-\hat{y}(t)}{\hat{y}(t)}",
    "eq_combined": r"p^{\mathrm{all}}(t)=\frac{\sum_{j=1}^{2} y^j(t)-"
    r"\sum_{j=1}^{2}\hat{y}^j(t)}{\sum_{j=1}^{2}\hat{y}^j(t)}",
    "eq_pooled": r"p_{\mathrm{pooled}}(t)=\sum_{k=1}^{K}\pi_k p_k(t)",
}


CORRECTED_FORMULAS = {
    "eq_model": r"\begin{aligned} y(t_i) &\overset{\mathrm{ind}}{\sim} "
    r"\operatorname{Poisson}(\mu_i),\quad i=1,\ldots,n, \\ "
    r"\log \mu_i &= \beta_0+f_{\mathrm{tr}}(t_i)+f_s(t_i)+\epsilon_i+o_i, \\ "
    r"f_{\mathrm{tr}} &\sim \operatorname{IWP}_2(\sigma_{\mathrm{tr}}),\quad "
    r"f_s \sim \operatorname{sGP}_1(\sigma_s;m=4),\quad "
    r"\epsilon_i \overset{\mathrm{iid}}{\sim} N(0,\sigma_\epsilon^2). "
    r"\end{aligned}",
    "eq_combined": r"p^{\mathrm{combined}}(t)=\frac{\sum_{j=1}^{J} y^j(t)-"
    r"\sum_{j=1}^{J}\hat{y}^j(t)}{\sum_{j=1}^{J}\hat{y}^j(t)}",
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def set_run_font(run, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")


def format_paragraph(paragraph, *, heading: bool = False) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(6 if heading else 0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.keep_with_next = heading


def add_text(document: Document, text: str, *, heading: bool = False):
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, heading=heading)
    run = paragraph.add_run(text)
    set_run_font(run, bold=heading)
    return paragraph


def latex_to_omml(latex: str, workdir: Path):
    key = hashlib.sha256(latex.encode("utf-8")).hexdigest()[:16]
    markdown_path = workdir / f"math-{key}.md"
    docx_path = workdir / f"math-{key}.docx"
    markdown_path.write_text(f"$$\n{latex}\n$$\n", encoding="utf-8")
    subprocess.run(
        ["pandoc", str(markdown_path), "-o", str(docx_path)],
        check=True,
        capture_output=True,
        text=True,
    )
    with zipfile.ZipFile(docx_path) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    math_nodes = root.xpath(".//m:oMath", namespaces=NS)
    if len(math_nodes) != 1:
        raise RuntimeError(f"Expected one OMML equation; found {len(math_nodes)}")
    return copy.deepcopy(math_nodes[0])


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(1, start)
    paragraph._p.append(end)


def add_equation(
    document: Document,
    latex: str,
    number: int,
    bookmark: str,
    bookmark_id: int,
    workdir: Path,
):
    paragraph = document.add_paragraph()
    paragraph.style = "Normal"
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_together = True
    tabs = paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    tabs.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left_tab = paragraph.add_run("\t")
    set_run_font(left_tab)
    paragraph._p.append(latex_to_omml(latex, workdir))
    number_run = paragraph.add_run(f"\t({number})")
    set_run_font(number_run)
    add_bookmark(paragraph, bookmark, bookmark_id)
    return paragraph


def clear_document_body(document: Document) -> None:
    body = document._element.body
    section_properties = body.sectPr
    for child in list(body):
        if child is not section_properties:
            body.remove(child)


def configure_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    for container in (section.header, section.footer):
        for paragraph in container.paragraphs:
            paragraph.text = ""


def build_baseline(reference: Path, output: Path, workdir: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(reference, output)
    document = Document(output)
    clear_document_body(document)
    configure_document(document)
    document.core_properties.title = "Appendices"
    document.core_properties.subject = "Appendix A: Model; Appendix B: P-score definition"
    document.core_properties.comments = (
        "Faithful Word conversion of Appendix A and Appendix B from appendix.tex."
    )

    add_text(document, "Appendices", heading=True)
    add_text(document, "Appendix A: Model", heading=True)
    add_text(document, SOURCE_PARAGRAPHS["model_intro"])
    add_equation(document, BASELINE_FORMULAS["eq_model"], 1, "eq_model", 1, workdir)
    add_text(document, SOURCE_PARAGRAPHS["model_explanation"])
    add_text(document, SOURCE_PARAGRAPHS["us_offset"])
    add_text(document, SOURCE_PARAGRAPHS["prediction"])
    add_equation(document, BASELINE_FORMULAS["eq_pscore"], 2, "eq_pscore", 2, workdir)
    add_text(document, SOURCE_PARAGRAPHS["combined_intro"])
    add_equation(
        document,
        BASELINE_FORMULAS["eq_combined"],
        3,
        "eq_combined",
        3,
        workdir,
    )
    add_text(document, SOURCE_PARAGRAPHS["combined_where"])
    add_text(document, SOURCE_PARAGRAPHS["pooled_intro"])
    add_equation(document, BASELINE_FORMULAS["eq_pooled"], 4, "eq_pooled", 4, workdir)
    add_text(document, SOURCE_PARAGRAPHS["pooled_where"])
    add_text(document, "Appendix B: P-score definition", heading=True)
    add_text(document, SOURCE_PARAGRAPHS["b_definition"])
    add_text(document, SOURCE_PARAGRAPHS["b_combined"])
    add_text(document, SOURCE_PARAGRAPHS["b_pooled"])

    document.save(output)


def make_run(text: str, *, deleted: bool = False):
    run = etree.Element(f"{{{W_NS}}}r")
    properties = etree.SubElement(run, f"{{{W_NS}}}rPr")
    fonts = etree.SubElement(properties, f"{{{W_NS}}}rFonts")
    fonts.set(f"{{{W_NS}}}ascii", "Calibri")
    fonts.set(f"{{{W_NS}}}hAnsi", "Calibri")
    size = etree.SubElement(properties, f"{{{W_NS}}}sz")
    size.set(f"{{{W_NS}}}val", "22")
    tag = "delText" if deleted else "t"
    text_element = etree.SubElement(run, f"{{{W_NS}}}{tag}")
    if text.startswith(" ") or text.endswith(" "):
        text_element.set(f"{{{XML_NS}}}space", "preserve")
    text_element.text = text
    return run


def make_tab_run(text: str | None = None):
    """Create a Word run with a real tab element, optionally followed by text."""
    run = make_run(text or "")
    text_element = run.find(f"{{{W_NS}}}t")
    tab = etree.Element(f"{{{W_NS}}}tab")
    text_element.addprevious(tab)
    if text is None:
        run.remove(text_element)
    return run


def revision_element(kind: str, revision_id: int, author: str, date: str):
    element = etree.Element(f"{{{W_NS}}}{kind}")
    element.set(f"{{{W_NS}}}id", str(revision_id))
    element.set(f"{{{W_NS}}}author", author)
    element.set(f"{{{W_NS}}}date", date)
    return element


def paragraph_visible_text(paragraph) -> str:
    return "".join(paragraph.xpath(".//w:t/text()", namespaces=NS))


def replace_paragraph_tracked(
    root,
    old_text: str,
    new_text: str,
    next_revision_id: int,
    author: str,
    date: str,
) -> int:
    matches = [
        paragraph
        for paragraph in root.xpath(".//w:body/w:p", namespaces=NS)
        if paragraph_visible_text(paragraph) == old_text
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected exactly one paragraph for tracked replacement; found {len(matches)}: "
            f"{old_text[:80]}"
        )
    paragraph = matches[0]
    paragraph_properties = paragraph.find(f"{{{W_NS}}}pPr")
    for child in list(paragraph):
        if child is not paragraph_properties:
            paragraph.remove(child)
    deletion = revision_element("del", next_revision_id, author, date)
    deletion.append(make_run(old_text, deleted=True))
    insertion = revision_element("ins", next_revision_id + 1, author, date)
    insertion.append(make_run(new_text, deleted=False))
    paragraph.append(deletion)
    paragraph.append(insertion)
    return next_revision_id + 2


def find_bookmarked_paragraph(root, bookmark_name: str):
    bookmarks = root.xpath(
        f'.//w:bookmarkStart[@w:name="{bookmark_name}"]', namespaces=NS
    )
    if len(bookmarks) != 1:
        raise RuntimeError(f"Expected one bookmark named {bookmark_name}")
    return bookmarks[0].getparent()


def build_equation_paragraph(
    ppr,
    math_node,
    number: int,
    revision_id: int,
    author: str,
    date: str,
):
    paragraph = etree.Element(f"{{{W_NS}}}p")
    if ppr is not None:
        paragraph.append(copy.deepcopy(ppr))
    # Keep the alignment tabs and unchanged equation number outside the tracked
    # insertion. This lets Word apply the paragraph's center/right tab stops
    # while tracking only the revised mathematical content.
    paragraph.append(make_tab_run())
    insertion = revision_element("ins", revision_id, author, date)
    insertion.append(copy.deepcopy(math_node))
    paragraph.append(insertion)
    paragraph.append(make_tab_run(f"({number})"))
    return paragraph


def replace_equation_tracked(
    root,
    bookmark_name: str,
    corrected_math,
    number: int,
    next_revision_id: int,
    author: str,
    date: str,
) -> int:
    paragraph = find_bookmarked_paragraph(root, bookmark_name)
    ppr = paragraph.find(f"{{{W_NS}}}pPr")
    content = [child for child in list(paragraph) if child is not ppr]
    for child in content:
        paragraph.remove(child)
    deletion = revision_element("del", next_revision_id, author, date)
    for child in content:
        if child.tag not in {
            f"{{{W_NS}}}bookmarkStart",
            f"{{{W_NS}}}bookmarkEnd",
        }:
            deletion.append(child)
    paragraph.append(deletion)
    inserted_paragraph = build_equation_paragraph(
        ppr,
        corrected_math,
        number,
        next_revision_id + 1,
        author,
        date,
    )
    paragraph.addnext(inserted_paragraph)
    return next_revision_id + 2


def enable_track_revisions(settings_root) -> None:
    existing = settings_root.xpath(".//w:trackRevisions", namespaces=NS)
    if existing:
        return
    element = etree.Element(f"{{{W_NS}}}trackRevisions")
    settings_root.insert(0, element)


def write_docx_parts(source: Path, output: Path, replacements: dict[str, bytes]) -> None:
    with zipfile.ZipFile(source, "r") as input_archive, zipfile.ZipFile(
        output, "w", compression=zipfile.ZIP_DEFLATED
    ) as output_archive:
        for item in input_archive.infolist():
            data = replacements.get(item.filename, input_archive.read(item.filename))
            output_archive.writestr(item, data)


def build_tracked(baseline: Path, output: Path, workdir: Path) -> None:
    with zipfile.ZipFile(baseline) as archive:
        document_root = etree.fromstring(archive.read("word/document.xml"))
        settings_root = etree.fromstring(archive.read("word/settings.xml"))

    revision_id = 1
    author = "Codex"
    date = dt.datetime.now(dt.timezone.utc).replace(microsecond=0).isoformat().replace(
        "+00:00", "Z"
    )
    for key in (
        "model_intro",
        "model_explanation",
        "us_offset",
        "combined_intro",
        "combined_where",
        "b_combined",
        "b_pooled",
    ):
        revision_id = replace_paragraph_tracked(
            document_root,
            SOURCE_PARAGRAPHS[key],
            CORRECTED_PARAGRAPHS[key],
            revision_id,
            author,
            date,
        )

    revision_id = replace_equation_tracked(
        document_root,
        "eq_model",
        latex_to_omml(CORRECTED_FORMULAS["eq_model"], workdir),
        1,
        revision_id,
        author,
        date,
    )
    revision_id = replace_equation_tracked(
        document_root,
        "eq_combined",
        latex_to_omml(CORRECTED_FORMULAS["eq_combined"], workdir),
        3,
        revision_id,
        author,
        date,
    )
    enable_track_revisions(settings_root)
    output.parent.mkdir(parents=True, exist_ok=True)
    write_docx_parts(
        baseline,
        output,
        {
            "word/document.xml": etree.tostring(
                document_root, xml_declaration=True, encoding="UTF-8", standalone=True
            ),
            "word/settings.xml": etree.tostring(
                settings_root, xml_declaration=True, encoding="UTF-8", standalone=True
            ),
        },
    )


def validate_docx(path: Path, *, tracked: bool) -> dict[str, object]:
    with zipfile.ZipFile(path) as archive:
        document_root = etree.fromstring(archive.read("word/document.xml"))
        settings_root = etree.fromstring(archive.read("word/settings.xml"))
    insertions = document_root.xpath(".//w:ins", namespaces=NS)
    deletions = document_root.xpath(".//w:del", namespaces=NS)
    equations = document_root.xpath(".//m:oMath", namespaces=NS)
    tracking_enabled = bool(
        settings_root.xpath(".//w:trackRevisions", namespaces=NS)
    )
    if tracked:
        if len(insertions) != 9 or len(deletions) != 9:
            raise RuntimeError(
                f"Tracked copy must contain 9 insertions and 9 deletions; found "
                f"{len(insertions)} and {len(deletions)}"
            )
        if not tracking_enabled:
            raise RuntimeError("Tracked copy does not enable trackRevisions")
        revisions = insertions + deletions
        if {item.get(f"{{{W_NS}}}author") for item in revisions} != {"Codex"}:
            raise RuntimeError("Tracked revisions have an unexpected author")
        identifiers = [item.get(f"{{{W_NS}}}id") for item in revisions]
        if len(identifiers) != len(set(identifiers)):
            raise RuntimeError("Tracked revision IDs are not unique")
        if len(equations) != 6:
            raise RuntimeError(f"Tracked copy should retain six OMML equation nodes; found {len(equations)}")
    else:
        if insertions or deletions or tracking_enabled:
            raise RuntimeError("Baseline unexpectedly contains tracked revisions")
        if len(equations) != 4:
            raise RuntimeError(f"Baseline must contain four OMML equations; found {len(equations)}")

    document = Document(path)
    section = document.sections[0]
    expected_inches = {
        "page_width": 8.5,
        "page_height": 11.0,
        "top_margin": 1.0,
        "bottom_margin": 1.0,
        "left_margin": 1.0,
        "right_margin": 1.0,
    }
    for attribute, expected in expected_inches.items():
        observed = getattr(section, attribute).inches
        if abs(observed - expected) > 0.01:
            raise RuntimeError(f"Unexpected {attribute}: {observed}")
    return {
        "path": str(path),
        "sha256": sha256(path),
        "bytes": path.stat().st_size,
        "insertions": len(insertions),
        "deletions": len(deletions),
        "equations": len(equations),
        "tracking_enabled": tracking_enabled,
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--reference", required=True, type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    parser.add_argument("--work-dir", type=Path)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    reference = args.reference.resolve()
    expected_reference_hash = (
        "6074ba7e00357b1ceab136ca137e07d710df41a8995b814c4781905ec4890a90"
    )
    if sha256(reference) != expected_reference_hash:
        raise RuntimeError("The manuscript formatting reference changed after distillation")
    args.output_dir.mkdir(parents=True, exist_ok=True)
    baseline = args.output_dir / "appendices_word_baseline.docx"
    tracked = args.output_dir / "appendices_word_tracked.docx"
    if args.work_dir is None:
        temporary = tempfile.TemporaryDirectory(prefix="appendices-word-")
        workdir = Path(temporary.name)
    else:
        temporary = None
        workdir = args.work_dir
        workdir.mkdir(parents=True, exist_ok=True)
    build_baseline(reference, baseline, workdir)
    build_tracked(baseline, tracked, workdir)
    baseline_validation = validate_docx(baseline, tracked=False)
    tracked_validation = validate_docx(tracked, tracked=True)
    print(baseline_validation)
    print(tracked_validation)
    if temporary is not None:
        temporary.cleanup()


if __name__ == "__main__":
    main()
